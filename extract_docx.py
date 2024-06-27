from docx import Document
from docx.shared import RGBColor, Pt
import os
from docx.oxml import OxmlElement

def extract_from_docx(docx_path, output_dir):
    doc = Document(docx_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    images_dir = os.path.join(output_dir, 'images')
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)

    text_output = []
    image_count = 1

    
    with open(os.path.join(output_dir, 'text.txt'), 'w', encoding='utf-8') as f:
        for paragraph in doc.paragraphs:
            text_output.append(paragraph.text)
            f.write(paragraph.text + '\n')

        f.write('\n')

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = ''
                    for paragraph in cell.paragraphs:
                        text += paragraph.text + '\n'
                    text_output.append(text.strip())
                    f.write(text.strip() + '\n')
            f.write('\n')

def extract_docx_text_details(docx_path, output_file):
    doc = Document(docx_path)
    with open(os.path.join(output_file, "detail.txt"), 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            text_content = para.text.strip()

            font_name = para.style.font.name if para.style.font and para.style.font.name else "Not specified"
            font_size = para.style.font.size.pt if para.style.font and para.style.font.size else "Not specified"

            # Check for bold, italic, underline
            runs = para.runs
            is_bold = any(run.bold for run in runs)
            is_italic = any(run.italic for run in runs)
            is_underline = any(run.underline for run in runs)

            # Text color (RGB)
            text_color = None
            if runs:
                first_run_color = runs[0].font.color
                if first_run_color and isinstance(first_run_color, RGBColor):
                    text_color = first_run_color.rgb

            # Convert RGB color to hex format if needed
            text_color_hex = "#{:02x}{:02x}{:02x}".format(text_color[0], text_color[1], text_color[2]) if text_color else "Not specified"

            # Write extracted details to file
            f.write(f"Text Content: {text_content}\n")
            f.write(f"Font Name: {font_name}\n")
            f.write(f"Font Size: {font_size}\n")
            f.write(f"Is Bold: {is_bold}\n")
            f.write(f"Is Italic: {is_italic}\n")
            f.write(f"Is Underline: {is_underline}\n")
            f.write(f"Text Color (RGB): {text_color}\n")
            f.write(f"Text Color (Hex): {text_color_hex}\n")
            f.write("-------------------------\n")
    print(f"text details saved to {output_file}")

def extract_and_uppercase_docx(input_path, output_path):
    output_path =  os.path.join(output_path, "uppercase.docx")
    doc = Document(input_path)
    new_doc = Document()

    # Copy paragraphs
    for para in doc.paragraphs:
        new_para = new_doc.add_paragraph()
        for run in para.runs:
            new_run = new_para.add_run(run.text.upper())
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            if run.font.color.rgb:
                new_run.font.color.rgb = RGBColor(run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
        new_para.alignment = para.alignment

    # Copy tables
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                new_cell.text = cell.text.upper()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        new_run = new_cell.paragraphs[-1].add_run(run.text.upper())
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                        if run.font.color.rgb:
                            new_run.font.color.rgb = RGBColor(run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])

    new_doc.save(output_path)
    print(f"Uppercased DOCX saved to {output_path}")